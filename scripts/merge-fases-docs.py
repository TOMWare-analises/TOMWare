# -*- coding: utf-8 -*-
"""Consolida documentos de fases da raiz do workspace em um unico arquivo."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

WORKSPACE_ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = Path(__file__).resolve().parents[1] / "docs"

SECTIONS = [
    ("Diagnostico inicial", WORKSPACE_ROOT / "Diagnóstico.docx"),
    ("Fase 1 — Estabilizacao e base operacional", WORKSPACE_ROOT / "fase1_ok.docx"),
    ("Fase 2 — AntiDebugMask (-dd)", WORKSPACE_ROOT / "fase2_ok.docx"),
    ("Fase 3 — ProcessEnumMask (-dp) e assinaturas (-sf)", WORKSPACE_ROOT / "fase3_ok.docx"),
    ("Fase 4 — Benchmark e corpus DBI-Log", WORKSPACE_ROOT / "fase4_ok.docx"),
    ("Fase 5 — Compilacao e testes locais", WORKSPACE_ROOT / "fase5_compilar_testar.docx"),
    ("Fase 6 — Testes na VM VMware", WORKSPACE_ROOT / "fase6_rodar_testes_na_VMWare.docx"),
]


def extract_paragraphs(path: Path) -> list[str]:
    doc = Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return lines


def extract_tables(path: Path) -> list[list[list[str]]]:
    doc = Document(str(path))
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
    return tables


def slugify(title: str) -> str:
    t = title.lower()
    t = re.sub(r"[^\w\s-]", "", t, flags=re.UNICODE)
    t = re.sub(r"\s+", "-", t.strip())
    return t[:60]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, lines: list[str]) -> None:
    for line in lines:
        doc.add_paragraph(line)


def add_table_markdown(lines: list[str], rows: list[list[str]]) -> None:
    if not rows:
        return
    header = rows[0]
    lines.append("")
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")
    lines.append("")


def build_markdown(sections: list[tuple[str, list[str], list[list[list[str]]]]]) -> str:
    now = datetime.now().strftime("%Y-%m-%d")
    md: list[str] = [
        "# TOMWare — Evolucao das melhorias por fases",
        "",
        f"Documento consolidado em {now} a partir dos registros de diagnostico e fases 1–6.",
        "",
        "Repositorio: [TOMWare-DBI/TOMWare](https://github.com/TOMWare-DBI/TOMWare)",
        "",
        "---",
        "",
    ]

    for idx, (title, paras, tables) in enumerate(sections, start=1):
        md.append(f"## {idx}. {title}")
        md.append("")
        for p in paras:
            md.append(p)
            md.append("")
        for table in tables:
            add_table_markdown(md, table)
        md.append("---")
        md.append("")

    md.append("## Resumo da evolucao")
    md.append("")
    md.append("| Fase | Foco principal | Entregavel |")
    md.append("|------|----------------|------------|")
    md.append("| Diagnostico | Lacunas vs DBI-Log e anti-debug | Plano de melhorias |")
    md.append("| 1 | Estabilidade, logging, `-da`/`-go`, scripts | Base operacional Release |")
    md.append("| 2 | AntiDebugMask (`-dd`) | PoC TestAntiDebug |")
    md.append("| 3 | ProcessEnumMask (`-dp`), assinaturas (`-sf`) | PoC TestProcessEnum |")
    md.append("| 4 | Benchmark, manifesto corpus, metricas | Scripts Fase 4 |")
    md.append("| 5 | Compilacao e validacao local | Build Release x64 |")
    md.append("| 6 | Execucao na VM, malware real, evidencias | Avaliacao em amostras reais |")
    md.append("")

    return "\n".join(md)


def build_docx(sections: list[tuple[str, list[str], list[list[list[str]]]]]) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "TOMWare — Evolucao das melhorias por fases", level=0)
    doc.add_paragraph(
        f"Documento consolidado em {datetime.now():%d/%m/%Y} "
        "a partir do diagnostico inicial e dos registros das fases 1 a 6."
    )
    doc.add_paragraph("Repositorio: https://github.com/TOMWare-DBI/TOMWare")

    for idx, (title, paras, tables) in enumerate(sections, start=1):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        add_heading(doc, f"{idx}. {title}", level=1)
        add_paragraphs(doc, paras)
        for table in tables:
            if not table:
                continue
            t = doc.add_table(rows=len(table), cols=len(table[0]))
            t.style = "Table Grid"
            for r_idx, row in enumerate(table):
                for c_idx, cell in enumerate(row):
                    t.rows[r_idx].cells[c_idx].text = cell
            doc.add_paragraph("")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "Resumo da evolucao", level=1)
    summary = [
        ("Diagnostico", "Lacunas vs DBI-Log e anti-debug", "Plano de melhorias"),
        ("Fase 1", "Estabilidade, logging, scripts", "Base operacional Release"),
        ("Fase 2", "AntiDebugMask (-dd)", "PoC TestAntiDebug"),
        ("Fase 3", "ProcessEnumMask (-dp), -sf", "PoC TestProcessEnum"),
        ("Fase 4", "Benchmark e corpus", "Scripts Fase 4"),
        ("Fase 5", "Compilacao local", "Build Release x64"),
        ("Fase 6", "VM e malware real", "Avaliacao em amostras reais"),
    ]
    st = doc.add_table(rows=1 + len(summary), cols=3)
    st.style = "Table Grid"
    hdr = st.rows[0].cells
    hdr[0].text = "Fase"
    hdr[1].text = "Foco principal"
    hdr[2].text = "Entregavel"
    for i, row in enumerate(summary, start=1):
        st.rows[i].cells[0].text = row[0]
        st.rows[i].cells[1].text = row[1]
        st.rows[i].cells[2].text = row[2]

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    collected: list[tuple[str, list[str], list[list[list[str]]]]] = []

    for title, path in SECTIONS:
        if not path.exists():
            raise FileNotFoundError(f"Arquivo nao encontrado: {path}")
        collected.append((title, extract_paragraphs(path), extract_tables(path)))

    md_path = OUT_DIR / "EVOLUCAO-MELHORIAS-POR-FASES.md"
    docx_path = OUT_DIR / "EVOLUCAO-MELHORIAS-POR-FASES.docx"

    md_path.write_text(build_markdown(collected), encoding="utf-8")
    build_docx(collected).save(str(docx_path))

    print(f"OK: {md_path}")
    print(f"OK: {docx_path}")


if __name__ == "__main__":
    main()
